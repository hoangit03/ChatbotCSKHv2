"""
app/infrastructure/parser/extractors/docx_parser.py
"""
from __future__ import annotations
import io, re
from app.core.interfaces.parser_port import ParsedChunk, ParsedDocument, ParserPort

class DocxParser(ParserPort):

    @property
    def supported_extensions(self) -> frozenset[str]:
        return frozenset({".docx"})

    async def parse(self, content: bytes, file_name: str) -> ParsedDocument:
        from docx import Document
        doc = Document(io.BytesIO(content))
        
        parts: list[str] = []
        tables: list[list[list[str]]] = []

        # 1. Đọc paragraphs chuẩn (bao gồm cả nội dung trong các khối thông thường)
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # 2. Đọc tables
        for table in doc.tables:
            tbl_data: list[list[str]] = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                tbl_data.append(row_data)
            
            if tbl_data:
                tables.append(tbl_data)
                parts.append(self._tbl_to_text(tbl_data))

        full_text = "\n\n".join(parts)
        
        if not full_text.strip():
            # Fallback: Nếu cách đọc trên không ra text, thử đọc thô xml (cho các trường hợp đặc biệt)
            full_text = self._fallback_parse(doc)

        return ParsedDocument(
            file_name=file_name,
            file_type="docx",
            chunks=[ParsedChunk(page=1, text=full_text, tables=tables)],
        )

    def _tbl_to_text(self, rows: list[list[str]]) -> str:
        if not rows: return ""
        header = rows[0]
        lines = []
        for row in rows[1:]:
            line = " | ".join(f"{h}: {v}" for h, v in zip(header, row) if v.strip())
            if line: lines.append(line)
        return "\n".join(lines)

    def _fallback_parse(self, doc) -> str:
        """Đọc thô xml nếu các methods trên không lấy được text."""
        parts = []
        for p in doc.element.xpath('//w:t'):
            if p.text: parts.append(p.text)
        return " ".join(parts)